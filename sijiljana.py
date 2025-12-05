# Streamlit App with Enhanced Design

import streamlit as st
import pandas as pd
import io
import zipfile
from PIL import Image, ImageDraw, ImageFont
from docx import Document

st.set_page_config(page_title="Jana Sijil Automatik", layout="wide")

# ================== HEADER (5,4,1) ==================
st.markdown("""
<div style='background: #f7faff; padding:25px; border-radius:15px; box-shadow:0 4px 10px rgba(0,0,0,0.08); border-bottom:4px solid #d0e3ff;'>
    <div style='display:flex; align-items:center; gap:20px;'>
        <span style='font-size:60px;'>📜</span>
        <div>
            <h1 style='margin:0; font-weight:700;'>Jana Sijil Automatik</h1>
            <p style='margin-top:6px; font-size:16px;'>Masukkan nama & IC murid, sistem akan jana sijil ikut template Word.</p>
        </div>
    </div>
</div>
<hr style='margin-top:25px;'>
""", unsafe_allow_html=True)

# ================== SESSION STATE ==================
if "df" not in st.session_state:
    st.session_state.df = pd.DataFrame({"Nama": [""], "IC": [""]})
if "generated" not in st.session_state:
    st.session_state.generated = False

# ================== UPLOAD TEMPLATE (10,11) ==================
st.markdown("""
<div style='padding:20px; border:2px dashed #0b5ed7; border-radius:12px; background:#eef5ff;'>
    <h3>📄 Upload Template Sijil (.docx)</h3>
</div>
""", unsafe_allow_html=True)

template_file = st.file_uploader("Muat naik template sijil Word (.docx)", type=["docx"])

st.markdown("""
<div style="border:2px solid #0b5ed7;background-color:#eef5ff;padding:12px;border-radius:10px; margin-top:10px;">
<b style='color:#0b5ed7;'>ℹ️ Nota Penting:</b><br>
Template mesti mengandungi placeholder:<br>
<code>{NAMA}</code> & <code>{IC}</code>
</div>
""", unsafe_allow_html=True)

# ================== TABLE INPUT (6,7,8,9) ==================
st.markdown("""
<div style='padding:20px; background:white; border-radius:12px; box-shadow:0 3px 12px rgba(0,0,0,0.08); margin-top:25px;'>
    <h3>👥 Senarai Nama & IC</h3>
    <p>Isi seperti jadual Excel. Boleh tambah baris bila perlu.</p>
</div>
""", unsafe_allow_html=True)

st.session_state.df = st.data_editor(
    st.session_state.df,
    num_rows="dynamic",
    use_container_width=True,
    hide_index=True
)

st.button("➕ Tambah Nama", key="addname", help="Tambah baris baru", type="primary")

# ================== FUNCTIONS ==================
def generate_docx(template_bytes, name, ic):
    doc = Document(io.BytesIO(template_bytes))
    for p in doc.paragraphs:
        p.text = p.text.replace("{NAMA}", name).replace("{IC}", ic)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace("{NAMA}", name).replace("{IC}", ic)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def docx_to_png(docx_bytes):
    doc = Document(io.BytesIO(docx_bytes.read()))
    section = doc.sections[0]
    width, height = section.page_width, section.page_height
    def emu_to_px(emu): return int(emu / 914400 * 300)
    W, H = emu_to_px(width), emu_to_px(height)
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    try: font = ImageFont.truetype("DejaVuSans-Bold.ttf", 72)
    except: font = ImageFont.load_default()
    y = 300
    for p in doc.paragraphs:
        if p.text.strip():
            draw.text((W//2, y), p.text, anchor="mm", fill="black", font=font)
            y += 150
    img_bytes = io.BytesIO()
    img.save(img_bytes, format="PNG")
    img_bytes.seek(0)
    return img_bytes

# ================== BUTTON JANA ==================
if template_file and st.button("🚀 Jana Sijil", type="primary"):
    st.session_state.generated = True
    st.session_state.template_bytes = template_file.read()

# ================== OUTPUT (13,14,15) ==================
if st.session_state.generated:
    st.subheader("📜 Senarai Sijil Murid")

    for idx, row in st.session_state.df.iterrows():
        with st.container():
            st.markdown("""
            <div style='padding:15px; background:white; border-radius:12px; box-shadow:0 3px 12px rgba(0,0,0,0.07); margin-bottom:15px;'>
            """, unsafe_allow_html=True)

            cols = st.columns([5, 2])
            name = row["Nama"]
            ic = row["IC"]

            cols[0].markdown(f"<h4 style='margin:0; font-weight:700;'>{name}</h4>", unsafe_allow_html=True)

            if cols[1].button("📥 Download", key=f"dl_{idx}"):
                docx_bytes = generate_docx(st.session_state.template_bytes, name, ic)
                png_bytes = docx_to_png(io.BytesIO(docx_bytes.getvalue()))
                cols[1].download_button("⬇️ Word", docx_bytes, f"Sijil_{name}.docx")
                cols[1].download_button("⬇️ PNG", png_bytes, f"Sijil_{name}.png")

# ================== DOWNLOAD SEMUA (18) ==================
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("📦 Muat Turun Semua", type="primary"):
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w") as zipf:
            for idx, row in st.session_state.df.iterrows():
                name = row["Nama"]
                ic = row["IC"]
                docx_bytes = generate_docx(st.session_state.template_bytes, name, ic)
                png_bytes = docx_to_png(io.BytesIO(docx_bytes.getvalue()))
                zipf.writestr(f"Sijil_{name}.docx", docx_bytes.getvalue())
                zipf.writestr(f"Sijil_{name}.png", png_bytes.getvalue())
        zip_buf.seek(0)
        st.download_button("⬇️ Download ZIP", zip_buf, "Sijil_All.zip")

# ================== FOOTER (19) ==================
st.markdown("""
<div style="padding:15px; border-radius:10px; border:1px solid #ccc; background:#f9f9f9;">
<h3>Dibangunkan oleh</h3>
<ul>
<li>Narjihah binti Mohd Hashim</li>
<li>Mohamad Adham bin Abdul Malek</li>
<li>Nor Aqilah binti Aliasam</li>
<li>Muhammad Amin bin Muhammad Puzi</li>
<li>Siti Nur Amira binti Abdul Talib</li>
</ul>
<b>SMK Sultan Muzafar Shah 1</b><br>
<b>Tahun 2025</b>
</div>
""", unsafe_allow_html=True)